import os
from docx import Document
from docx.shared import Pt
from datetime import datetime

def create_work_report(full_name, date, work_plan, status_of_completion):
    # Create a new Document
    doc = Document()
    
    # Title
    doc.add_heading('Daily Work Report', 0)
    
    # Create a table with 2 columns and 4 rows
    table = doc.add_table(rows=4, cols=2)
    
    # Define column widths (if desired; adjust as needed)
    for row in table.rows:
        row.cells[0].width = Pt(100)  # This line sets the width for the first column
        row.cells[1].width = Pt(300)  # This line sets the width for the second column
    
    # Populate the table
    table.cell(0, 0).text = 'Full Name:'
    table.cell(0, 1).text = full_name
    
    table.cell(1, 0).text = 'Date:'
    table.cell(1, 1).text = date
    
    table.cell(2, 0).text = 'Work Plan:'
    table.cell(2, 1).text = work_plan
    
    table.cell(3, 0).text = 'Status of Completion:'
    table.cell(3, 1).text = status_of_completion
    
    # Save the document to Desktop
    desktop_path = os.path.expanduser("~/Desktop")
    file_name = f"{full_name.replace(' ', '_')}_Daily_Work_Report_{date.replace('/', '-')}.docx"
    full_path = os.path.join(desktop_path, file_name)
    doc.save(full_path)
    
    print(f"Report saved as {full_path}")

# Collect input from the user
full_name = input("Enter full name: ")

# Get the current date in the desired format
current_date = datetime.now().strftime("%d/%m/%Y")

print(f"Date set to: {current_date}")

print("Enter work plan (press Enter twice to end):")
work_plan = ""
while True:
    line = input()
    if line == "":
        break
    work_plan += line + "\n"

print("Enter status of completion (press Enter twice to end):")
status_of_completion = ""
while True:
    line = input()
    if line == "":
        break
    status_of_completion += line + "\n"

# Generate the report
create_work_report(full_name, current_date, work_plan.strip(), status_of_completion.strip())
